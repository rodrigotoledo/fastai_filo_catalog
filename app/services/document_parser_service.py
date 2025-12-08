import os
import tempfile
from typing import Dict, List, Optional, Any
from pathlib import Path
import logging

# PDF parsing
import PyPDF2
import pdfplumber

# DOCX parsing
from docx import Document

# OCR for images
import pytesseract
from PIL import Image

# CSV/Excel parsing
import pandas as pd

# Markdown parsing (simple text extraction)
import re

from app.services.ai_service import AIService

logger = logging.getLogger(__name__)

class DocumentParserService:
    """
    Service for parsing various document formats to extract client information.
    Uses AI to intelligently extract structured data from unstructured text.
    """

    def __init__(self, ai_service: AIService):
        self.ai_service = ai_service

    def parse_document(self, file_path: str, filename: str, extraction_prompt: Optional[str] = None) -> Dict[str, Any]:
        """
        Parse a document and extract client information using AI.

        Args:
            file_path: Path to the uploaded file
            filename: Original filename with extension

        Returns:
            Dict containing extracted client data
        """
        file_extension = Path(filename).suffix.lower()

        # Extract text based on file type
        text_content = self._extract_text(file_path, file_extension)

        if not text_content:
            raise ValueError(f"Could not extract text from {filename}")

        # Use AI to extract structured client data
        return self._extract_client_data_with_ai(text_content, filename, extraction_prompt)

    def parse_document_from_bytes(self, file_content: bytes, filename: str, extraction_prompt: Optional[str] = None) -> Dict[str, Any]:
        """
        Parse a document from bytes content and extract client information using AI.
        Avoids creating temporary files when possible.

        Args:
            file_content: File content as bytes
            filename: Original filename with extension

        Returns:
            Dict containing extracted client data
        """
        file_extension = Path(filename).suffix.lower()

        # Extract text based on file type - try to work in memory when possible
        text_content = self._extract_text_from_bytes(file_content, file_extension, filename)

        if not text_content:
            raise ValueError(f"Could not extract text from {filename}")

        # Use AI to extract structured client data
        return self._extract_client_data_with_ai(text_content, filename, extraction_prompt)

    def generate_client_embedding(self, extracted_data: Dict[str, Any]) -> Optional[List[float]]:
        try:
            # Monta texto rico do cliente
            client_text_parts = []
            if extracted_data.get('name'):
                client_text_parts.append(extracted_data['name'])
            if extracted_data.get('cpf'):
                client_text_parts.append(f"CPF {extracted_data['cpf']}")
            if extracted_data.get('email'):
                client_text_parts.append(extracted_data['email'])
            if extracted_data.get('phone'):
                client_text_parts.append(extracted_data['phone'])

            address = extracted_data.get('address', {})
            addr_parts = [address.get(k) for k in ['street', 'city', 'state'] if address.get(k)]
            if addr_parts:
                client_text_parts.append("Endereço: " + ", ".join(addr_parts))

            client_text = ". ".join(client_text_parts)
            if not client_text.strip():
                return None

            # AQUI É O PULO DO GATO:
            # MUDANÇA CRÍTICA
            embedding = self.ai_service.generate_clip_text_embedding(client_text)
            return embedding
        except Exception as e:
            logger.error(f"Error generating client embedding: {str(e)}")
            return None

    def _extract_text(self, file_path: str, file_extension: str) -> str:
        """
        Extract text from various file formats.

        Args:
            file_path: Path to the file
            file_extension: File extension (e.g., '.pdf', '.docx')

        Returns:
            Extracted text content
        """
        try:
            if file_extension == '.pdf':
                return self._extract_pdf_text(file_path)
            elif file_extension == '.docx':
                return self._extract_docx_text(file_path)
            elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                return self._extract_image_text(file_path)
            elif file_extension == '.csv':
                return self._extract_csv_text(file_path)
            elif file_extension in ['.xlsx', '.xls']:
                return self._extract_excel_text(file_path)
            elif file_extension == '.md':
                return self._extract_markdown_text(file_path)
            elif file_extension == '.txt':
                return self._extract_plain_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise

    def _extract_text_from_bytes(self, file_content: bytes, file_extension: str, filename: str) -> str:
        """
        Extract text from file content in bytes. Creates temporary file only when necessary.

        Args:
            file_content: File content as bytes
            file_extension: File extension (e.g., '.pdf', '.docx')
            filename: Original filename for temporary file creation

        Returns:
            Extracted text content
        """
        try:
            # For text-based files, work directly in memory
            if file_extension in ['.md', '.txt', '.csv']:
                if file_extension == '.csv':
                    return self._extract_csv_text_from_bytes(file_content)
                elif file_extension == '.md':
                    return self._extract_markdown_text_from_bytes(file_content)
                elif file_extension == '.txt':
                    return self._extract_plain_text_from_bytes(file_content)

            # For binary files (PDF, DOCX, images, Excel), create temporary file
            else:
                # Create temporary file only when necessary
                import tempfile
                import os

                with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
                    temp_file.write(file_content)
                    temp_file_path = temp_file.name

                try:
                    if file_extension == '.pdf':
                        return self._extract_pdf_text(temp_file_path)
                    elif file_extension == '.docx':
                        return self._extract_docx_text(temp_file_path)
                    elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                        return self._extract_image_text(temp_file_path)
                    elif file_extension in ['.xlsx', '.xls']:
                        return self._extract_excel_text(temp_file_path)
                    else:
                        raise ValueError(f"Unsupported file type: {file_extension}")
                finally:
                    # Clean up temporary file
                    if os.path.exists(temp_file_path):
                        os.unlink(temp_file_path)

        except Exception as e:
            logger.error(f"Error extracting text from bytes for {filename}: {str(e)}")
            raise

    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF files using multiple methods."""
        text = ""

        # Try pdfplumber first (better for structured documents)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.warning(f"pdfplumber failed: {str(e)}")

        # Fallback to PyPDF2 if pdfplumber didn't work well
        if not text.strip():
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            except Exception as e:
                logger.warning(f"PyPDF2 failed: {str(e)}")

        return text.strip()

    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX files."""
        doc = Document(file_path)
        text = ""

        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"

        return text.strip()

    def _extract_image_text(self, file_path: str) -> str:
        """Extract text from images using OCR."""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang='por+eng')  # Portuguese + English
            return text.strip()
        except Exception as e:
            logger.error(f"OCR failed for {file_path}: {str(e)}")
            return ""

    def _extract_csv_text(self, file_path: str) -> str:
        """Extract text from CSV files."""
        try:
            df = pd.read_csv(file_path, encoding='utf-8')
            # Convert to readable text format
            text = f"CSV Data with {len(df)} rows and {len(df.columns)} columns:\n"
            text += "Columns: " + ", ".join(df.columns.tolist()) + "\n\n"

            # Add first few rows as examples
            text += "Sample data:\n"
            text += df.head(10).to_string(index=False)

            return text
        except Exception as e:
            logger.error(f"CSV parsing failed: {str(e)}")
            return ""

    def _extract_excel_text(self, file_path: str) -> str:
        """Extract text from Excel files."""
        try:
            df = pd.read_excel(file_path)
            # Convert to readable text format
            text = f"Excel Data with {len(df)} rows and {len(df.columns)} columns:\n"
            text += "Columns: " + ", ".join(df.columns.tolist()) + "\n\n"

            # Add first few rows as examples
            text += "Sample data:\n"
            text += df.head(10).to_string(index=False)

            return text
        except Exception as e:
            logger.error(f"Excel parsing failed: {str(e)}")
            return ""

    def _extract_markdown_text(self, file_path: str) -> str:
        """Extract text from Markdown files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()

            # Remove markdown formatting for cleaner text
            # Remove headers
            content = re.sub(r'^#{1,6}\s+.*$', '', content, flags=re.MULTILINE)
            # Remove links
            content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', content)
            # Remove emphasis
            content = re.sub(r'\*\*([^\*]+)\*\*', r'\1', content)
            content = re.sub(r'\*([^\*]+)\*', r'\1', content)
            content = re.sub(r'_([^_]+)_', r'\1', content)

            return content.strip()
        except Exception as e:
            logger.error(f"Markdown parsing failed: {str(e)}")
            return ""

    def _extract_plain_text(self, file_path: str) -> str:
        """Extract text from plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            logger.error(f"Plain text parsing failed: {str(e)}")
            return ""

    def _extract_csv_text_from_bytes(self, file_content: bytes) -> str:
        """Extract text from CSV file content in bytes."""
        try:
            import io
            # Decode bytes to string
            content_str = file_content.decode('utf-8')
            # Create StringIO for pandas
            string_io = io.StringIO(content_str)
            df = pd.read_csv(string_io)

            # Convert to readable text format
            text = f"CSV Data with {len(df)} rows and {len(df.columns)} columns:\n"
            text += "Columns: " + ", ".join(df.columns.tolist()) + "\n\n"

            # Add first few rows as examples
            text += "Sample data:\n"
            text += df.head(10).to_string(index=False)

            return text
        except Exception as e:
            logger.error(f"CSV parsing from bytes failed: {str(e)}")
            return ""

    def _extract_markdown_text_from_bytes(self, file_content: bytes) -> str:
        """Extract text from Markdown file content in bytes."""
        try:
            # Decode bytes to string
            content = file_content.decode('utf-8')

            # Remove markdown formatting for cleaner text
            # Remove headers
            content = re.sub(r'^#{1,6}\s+.*$', '', content, flags=re.MULTILINE)
            # Remove links
            content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', content)
            # Remove emphasis
            content = re.sub(r'\*\*([^\*]+)\*\*', r'\1', content)
            content = re.sub(r'\*([^\*]+)\*', r'\1', content)
            content = re.sub(r'_([^_]+)_', r'\1', content)

            return content.strip()
        except Exception as e:
            logger.error(f"Markdown parsing from bytes failed: {str(e)}")
            return ""

    def _extract_plain_text_from_bytes(self, file_content: bytes) -> str:
        """Extract text from plain text file content in bytes."""
        try:
            # Decode bytes to string
            return file_content.decode('utf-8').strip()
        except Exception as e:
            logger.error(f"Plain text parsing from bytes failed: {str(e)}")
            return ""

    def _extract_client_data_with_ai(self, text_content: str, filename: str, extraction_prompt: Optional[str] = None) -> Dict[str, Any]:
        """
        Use AI to extract structured client data from text content.

        Args:
            text_content: Raw text extracted from document
            filename: Original filename for context

        Returns:
            Dict with extracted client information
        """
        prompt = extraction_prompt if extraction_prompt else f"""
        Analyze the following text extracted from a document ({filename}) and extract client information.
        Look for personal details that could be used to create a client record.

        Text content:
        {{text_content}}

        Please extract the following information if available:
        - Full name (nome completo)
        - CPF (Brazilian tax ID)
        - Email address
        - Phone number
        - Date of birth
        - Address information (street, number, city, state, postal code)
        - Any other relevant personal or business information

        Format your response as a JSON object with these possible keys:
        {{{{
            "name": "full name if found",
            "cpf": "CPF if found",
            "email": "email if found",
            "phone": "phone if found",
            "date_of_birth": "YYYY-MM-DD if found",
            "address": {{{{
                "street": "street address",
                "number": "number",
                "complement": "complement if any",
                "neighborhood": "neighborhood",
                "city": "city",
                "state": "state",
                "postal_code": "postal code"
            }}}},
            "notes": "any additional relevant information",
            "confidence": "high/medium/low based on how well the data matches"
        }}}}

        If information is not found, omit the key or set it to null.
        Be conservative - only extract information that clearly appears to be client data.
        """

        try:
            # Use AI service to process the text with custom prompt
            extracted_data = self.ai_service.process_text_with_custom_prompt(text_content, prompt)

            return extracted_data

        except Exception as e:
            logger.error(f"AI extraction failed: {str(e)}")
            return {"error": "Failed to extract data with AI", "raw_text": text_content[:500]}

    def _parse_ai_response(self, ai_response: str) -> Dict[str, Any]:
        """
        Parse the AI response into structured client data.
        This is a placeholder - actual implementation would depend on AI service output format.
        """
        # Placeholder implementation - in reality, this would parse JSON from AI response
        try:
            # Assuming AI returns JSON-like response
            import json
            return json.loads(ai_response)
        except:
            # Fallback: return basic structure
            return {
                "name": None,
                "cpf": None,
                "email": None,
                "phone": None,
                "date_of_birth": None,
                "address": {
                    "street": None,
                    "number": None,
                    "complement": None,
                    "neighborhood": None,
                    "city": None,
                    "state": None,
                    "postal_code": None
                },
                "notes": ai_response[:500] if ai_response else None,
                "confidence": "low"
            }

    def validate_extracted_data(self, data: Dict[str, Any]) -> List[str]:
        """
        Validate extracted client data and return list of validation errors.
        More permissive validation - only blocks on critical issues.
        """
        errors = []

        # Check required fields - only name is truly required
        if not data.get('name'):
            errors.append("Name is required")

        # For other fields, we'll be more permissive and just warn
        # CPF validation (optional)
        cpf = data.get('cpf')
        if cpf and not self._validate_cpf(cpf):
            # Instead of error, just remove invalid CPF
            data['cpf'] = None

        # Email validation (optional)
        email = data.get('email')
        if email and not self._validate_email(email):
            errors.append("Invalid email format")

        # Phone validation (optional)
        phone = data.get('phone')
        if phone and not self._validate_phone(phone):
            # Instead of error, just remove invalid phone
            data['phone'] = None

        return errors

    def _validate_cpf(self, cpf: str) -> bool:
        """Validate Brazilian CPF format (very relaxed for document extraction)."""
        # Remove non-numeric characters
        cpf_clean = re.sub(r'\D', '', cpf)

        # Accept any CPF-like string with at least 8 digits (be very permissive)
        return len(cpf_clean) >= 8

    def _validate_email(self, email: str) -> bool:
        """Validate email format (relaxed)."""
        # Very basic validation - just check for @ and .
        return '@' in email and '.' in email and len(email) > 5

    def _validate_phone(self, phone: str) -> bool:
        """Validate Brazilian phone format."""
        # Remove non-numeric characters
        phone = re.sub(r'\D', '', phone)

        # Accept Brazilian phones: 10 or 11 digits (with or without area code)
        # Also accept international format with country code (12-13 digits for Brazil)
        return len(phone) in [10, 11, 12, 13] and phone.startswith(('1', '2', '3', '4', '5', '6', '7', '8', '9'))
